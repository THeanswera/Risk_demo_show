"""
Генерация Word-отчёта (docx) с примером расчёта ИПР.
Требует: python-docx
"""

import io

import pandas as pd

from core.constants import R_NORM
from utils.helpers import safe_float


def generate_report_docx(
    df_scen: pd.DataFrame,
    df_grp: pd.DataFrame,
    df_scen_calc: pd.DataFrame,
    df_rows_calc: pd.DataFrame,
    df_agg: pd.DataFrame,
    r_total: float,
) -> bytes:
    """
    Генерирует Word-документ с полным примером расчёта ИПР.

    Параметры
    ----------
    df_scen       : исходные данные сценариев (из session_state)
    df_grp        : исходные данные групп (из session_state)
    df_scen_calc  : сценарии с расчётными коэффициентами
    df_rows_calc  : результаты расчёта по группам
    df_agg        : агрегирование по сценариям
    r_total       : расчётная величина ИПР

    Возвращает
    ----------
    bytes - содержимое .docx файла
    """
    try:
        from docx import Document
        from docx.shared import Pt
        from docx.enum.text import WD_ALIGN_PARAGRAPH
    except ImportError as exc:
        raise ImportError(
            "Для генерации отчёта установите python-docx: pip install python-docx"
        ) from exc

    doc = Document()

    # Настройка стилей
    style_normal = doc.styles["Normal"]
    style_normal.font.name = "Times New Roman"
    style_normal.font.size = Pt(12)

    def _add_heading(text: str, level: int = 1) -> None:
        h = doc.add_heading(text, level=level)
        h.runs[0].font.name = "Times New Roman"

    def _add_para(text: str, bold: bool = False) -> None:
        p = doc.add_paragraph()
        run = p.add_run(text)
        run.font.name = "Times New Roman"
        run.font.size = Pt(12)
        run.bold = bold

    def _add_table(headers: list[str], rows: list[list]) -> None:
        table = doc.add_table(rows=1 + len(rows), cols=len(headers))
        table.style = "Table Grid"
        hdr_cells = table.rows[0].cells
        for i, h in enumerate(headers):
            hdr_cells[i].text = h
            for para in hdr_cells[i].paragraphs:
                for run in para.runs:
                    run.bold = True
                    run.font.name = "Times New Roman"
                    run.font.size = Pt(11)
        for row_data in rows:
            row_cells = table.add_row().cells
            for i, val in enumerate(row_data):
                row_cells[i].text = str(val)
                for para in row_cells[i].paragraphs:
                    for run in para.runs:
                        run.font.name = "Times New Roman"
                        run.font.size = Pt(11)

    # Титульный раздел
    title = doc.add_paragraph()
    title.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run_title = title.add_run(
        "РАСЧЁТ ИНДИВИДУАЛЬНОГО ПОЖАРНОГО РИСКА\n"
        "по Методике (Приказ МЧС России от 14.11.2022 № 1140)"
    )
    run_title.bold = True
    run_title.font.name = "Times New Roman"
    run_title.font.size = Pt(14)

    doc.add_paragraph()

    # Раздел 0 (опциональный): Расчёт времени эвакуации
    evac_result = None
    evac_group = None
    try:
        import streamlit as _st
        evac_result = _st.session_state.get("evac_result")
        evac_group = _st.session_state.get("evac_result_group")
    except Exception:
        pass

    if evac_result and evac_result.get("segments_detail"):
        _add_heading("Расчёт времени эвакуации (Приложение 6)", level=1)
        if evac_group:
            _add_para(f"Группа контингента: {evac_group}")

        seg_headers = ["N", "Тип", "Длина (м)", "Ширина (м)", "D (м2/м2)",
                       "V (м/мин)", "q (м/мин)", "t (мин)", "Скопление"]
        seg_rows = []
        for s in evac_result["segments_detail"]:
            seg_rows.append([
                s["N"], s["тип"],
                f"{s['длина']:.3f}", f"{s['ширина']:.3f}",
                f"{s['D']:.3f}", f"{s['V']:.3f}",
                f"{s['q']:.3f}", f"{s['t']:.3f}",
                "Да" if s["скопление"] else "Нет",
            ])
        _add_table(seg_headers, seg_rows)
        doc.add_paragraph()
        _add_para(f"tр = {evac_result['t_p']:.3f} мин", bold=True)
        _add_para(f"tск = {evac_result['t_ck']:.3f} мин", bold=True)
        doc.add_paragraph()

    # Раздел 1: Исходные данные
    _add_heading("1. Исходные данные", level=1)

    _add_heading("1.1. Сценарии пожара", level=2)
    scen_headers = [
        "Сценарий i", "Тип здания", "Qп (год⁻¹)", "tпр (ч/сут)",
        "tбл (мин)", "АУП", "ПС", "СОУЭ", "ПДЗ",
    ]
    scen_rows = []
    for _, r in df_scen.iterrows():
        # ИСПРАВЛЕНО: K_ап теперь бинарный (п. 15 Методики №1140)
        k_ap_val = safe_float(r.get("K_ап,i", 0))
        if "АУП соответствует? (K_ап=0.9)" in r.index:
            k_ap_display = "Да" if r.get("АУП соответствует? (K_ап=0.9)", False) else "Нет"
        else:
            k_ap_display = "Да" if k_ap_val >= 0.9 else "Нет"
        scen_rows.append([
            int(r.get("Сценарий i", 0)),
            r.get("Тип здания", ""),
            f"{safe_float(r.get('Q_п,i (год⁻¹)', 0)):.3e}",
            f"{safe_float(r.get('t_пр,i (ч/сут)', 0)):.1f}",
            f"{safe_float(r.get('t_бл,i (мин)', 0)):.2f}",
            k_ap_display,
            "Да" if r.get("ПС соответствует? (K_обн=0.8)", False) else "Нет",
            "Да" if r.get("СОУЭ соответствует? (K_СОУЭ=0.8)", False) else "Нет",
            "Да" if r.get("ПДЗ соответствует? (K_ПДЗ=0.8)", False) else "Нет",
        ])
    _add_table(scen_headers, scen_rows)

    doc.add_paragraph()
    _add_heading("1.2. Группы эвакуируемого контингента", level=2)
    grp_headers = ["ID", "Сценарий i", "Группа j", "tр (мин)", "tн.э (мин)", "tск (мин)"]
    grp_rows = []
    for _, r in df_grp.iterrows():
        grp_rows.append([
            int(r.get("ID", 0)),
            int(r.get("Сценарий i", 0)),
            r.get("Группа j", ""),
            f"{safe_float(r.get('t_р,i,j (мин)', 0)):.2f}",
            f"{safe_float(r.get('t_н.э,i,j (мин)', 0)):.2f}",
            f"{safe_float(r.get('t_ск,i,j (мин)', 0)):.2f}",
        ])
    _add_table(grp_headers, grp_rows)

    # Раздел 2: K_п.з
    doc.add_paragraph()
    _add_heading("2. Расчёт коэффициентов противопожарной защиты - формула (7)", level=1)

    kpz_rows = []
    for _, r in df_scen_calc.iterrows():
        si = int(r.get("Сценарий i", 0))
        k_obn = safe_float(r.get("K_обн,i", 0))
        k_soue = safe_float(r.get("K_СОУЭ,i", 0))
        k_pdz = safe_float(r.get("K_ПДЗ,i", 0))
        k_pz_val = safe_float(r.get("K_п.з,i", 0))

        k_obn_str = f"{k_obn:.1f} (ПС соответствует)" if k_obn > 0 else "0.0 (ПС не соответствует)"
        k_soue_str = f"{k_soue:.1f} (СОУЭ соответствует)" if k_soue > 0 else "0.0 (СОУЭ не соответствует)"
        k_pdz_str = f"{k_pdz:.1f} (ПДЗ соответствует)" if k_pdz > 0 else "0.0 (ПДЗ не соответствует)"

        _add_para(f"Сценарий {si}:")
        _add_para(f"  Kобн,{si} = {k_obn_str}")
        _add_para(f"  Kсоуэ,{si} = {k_soue_str}")
        _add_para(f"  Kпдз,{si} = {k_pdz_str}")
        formula_str = (
            f"  Kп.з,{si} = 1 - (1 - {k_obn:.1f}\u00d7{k_soue:.1f}) \u00d7 (1 - {k_obn:.1f}\u00d7{k_pdz:.1f}) "
            f"= 1 - (1 - {k_obn*k_soue:.4f}) \u00d7 (1 - {k_obn*k_pdz:.4f}) = {k_pz_val:.4f}"
        )
        _add_para(formula_str)
        kpz_rows.append([si, f"{k_obn:.1f}", f"{k_soue:.1f}", f"{k_pdz:.1f}", f"{k_pz_val:.4f}"])

    doc.add_paragraph()
    _add_table(["Сценарий i", "Kобн", "Kсоуэ", "Kпдз", "Kп.з"], kpz_rows)

    # Раздел 3: P_пр
    doc.add_paragraph()
    _add_heading("3. Расчёт вероятности присутствия - формула (5)", level=1)

    for _, r in df_scen_calc.iterrows():
        si = int(r.get("Сценарий i", 0))
        t_pr = safe_float(r.get("t_пр,i (ч/сут)", 0))
        p_pr = safe_float(r.get("P_пр,i", 0))
        _add_para(f"Pпр,{si} = tпр,{si} / 24 = {t_pr:.1f} / 24 = {p_pr:.4f}")

    # Раздел 4: P_э
    doc.add_paragraph()
    _add_heading("4. Расчёт вероятности эвакуации - формула (6)", level=1)

    for _, r in df_rows_calc.iterrows():
        row_id = int(r.get("ID", 0))
        si = int(r.get("Сценарий i", 0))
        grp = r.get("Группа j", "")
        t_bl = safe_float(r.get("t_бл,i (мин)", 0))
        t_p = safe_float(r.get("t_р,i,j (мин)", 0))
        t_ne = safe_float(r.get("t_н.э,i,j (мин)", 0))
        t_ck = safe_float(r.get("t_ск,i,j (мин)", 0))
        p_e = safe_float(r.get("P_э,i,j", 0))
        border = 0.8 * t_bl

        _add_para(f"ID {row_id} - Сценарий {si}, Группа \u00ab{grp}\u00bb:")
        _add_para(f"  0.8 \u00d7 tбл = 0.8 \u00d7 {t_bl:.2f} = {border:.2f} мин")
        _add_para(f"  tр + tн.э = {t_p:.2f} + {t_ne:.2f} = {t_p + t_ne:.2f} мин")
        _add_para(f"  tск = {t_ck:.2f} мин")

        # ИСПРАВЛЕНО: диагностика Pэ по п. 17 Методики №1140
        t_sum = t_p + t_ne
        if t_ck > 6.0:
            _add_para(f"  tск > 6 мин \u2192 Pэ = 0.000")
        elif t_sum < border:
            _add_para(
                f"  {t_sum:.2f} < {border:.2f} и tск = {t_ck:.2f} \u2264 6 \u2192 Pэ = 0.999"
            )
        elif t_sum >= t_bl:
            _add_para(f"  tр + tн.э = {t_sum:.2f} \u2265 tбл = {t_bl:.2f} \u2192 Pэ = 0.000")
        else:
            _add_para(
                f"  Промежуточная ветвь: "
                f"Pэ = 1 \u2212 (tр + tн.э) / tбл = 1 \u2212 {t_sum:.2f} / {t_bl:.2f} = {p_e:.4f}"
            )

    # Раздел 5: R_i,j
    doc.add_paragraph()
    _add_heading("5. Расчёт ИПР по группам - формула (4)", level=1)

    rij_table_rows = []
    for _, r in df_rows_calc.iterrows():
        row_id = int(r.get("ID", 0))
        si = int(r.get("Сценарий i", 0))
        grp = r.get("Группа j", "")
        q_p = safe_float(r.get("Q_п,i (год⁻¹)", 0))
        k_ap = safe_float(r.get("K_ап,i", 0))
        p_pr = safe_float(r.get("P_пр,i", 0))
        p_e = safe_float(r.get("P_э,i,j", 0))
        k_pz_val = safe_float(r.get("K_п.з,i", 0))
        r_val = safe_float(r.get("R_i,j", 0))

        _add_para(
            f"R{si},{row_id} = {q_p:.2e} \u00d7 (1 - {k_ap:.2f}) \u00d7 {p_pr:.4f} \u00d7 "
            f"(1 - {p_e:.4f}) \u00d7 (1 - {k_pz_val:.4f}) = {r_val:.3e}"
        )
        rij_table_rows.append([row_id, si, grp, f"{r_val:.3e}"])

    doc.add_paragraph()
    _add_table(["ID", "Сценарий i", "Группа j", "Ri,j"], rij_table_rows)

    # Раздел 6: R_i, R
    doc.add_paragraph()
    _add_heading("6. Определение расчётной величины ИПР - формулы (2)\u2013(3)", level=1)

    for _, r in df_agg.iterrows():
        si = r.get("Сценарий i", "")
        r_i = safe_float(r.get("R_i = max_j(R_i,j)", 0))
        if str(si) == "ИТОГО (R = max)":
            continue
        r_vals = df_rows_calc.loc[df_rows_calc["Сценарий i"].astype(str) == str(si), "R_i,j"]
        vals_str = ", ".join(f"{v:.3e}" for v in r_vals.tolist())
        _add_para(f"R{si} = max{{{vals_str}}} = {r_i:.3e} год\u207b\u00b9")

    _add_para(f"R = max{{Ri}} = {r_total:.3e} год\u207b\u00b9", bold=True)

    # Раздел 7: Заключение
    doc.add_paragraph()
    _add_heading("7. Заключение", level=1)

    passed = r_total <= R_NORM
    verdict = "ВЫПОЛНЕНО" if passed else "НЕ ВЫПОЛНЕНО"

    _add_para(f"R = {r_total:.3e} год\u207b\u00b9", bold=True)
    _add_para(f"Rнорм = {R_NORM:.1e} год\u207b\u00b9")
    _add_para(
        f"R {'\u2264' if passed else '>'} Rнорм - условие формулы (1) {verdict}",
        bold=True,
    )

    if passed:
        _add_para(
            "Расчётная величина индивидуального пожарного риска не превышает "
            "нормативного значения. Пожарный риск ДОПУСТИМЫЙ."
        )
    else:
        _add_para(
            "Расчётная величина индивидуального пожарного риска ПРЕВЫШАЕТ "
            "нормативное значение. Требуется разработка мероприятий по снижению пожарного риска."
        )

    # Сохранение в bytes
    buf = io.BytesIO()
    doc.save(buf)
    buf.seek(0)
    return buf.read()
